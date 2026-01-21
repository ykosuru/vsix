#!/usr/bin/env python3
"""
AstraCode Index Builder

Builds a searchable index from code files and documents (PDF, Excel, Word).
The index is saved as JSON and loaded by the AstraCode VS Code extension.

Usage:
    python astra-index.py /path/to/workspace
    python astra-index.py /path/to/workspace --output custom-index.json
    python astra-index.py folder1 folder2 folder3 --output index.json

Requirements:
    pip install pypdf2 openpyxl python-docx

Optional (for better PDF extraction):
    pip install pdfplumber
"""

import os
import sys
import json
import argparse
import re
from pathlib import Path
from datetime import datetime
from collections import defaultdict

# Optional imports - gracefully handle missing dependencies
PARSERS_AVAILABLE = {
    'pdf': False,
    'excel': False,
    'word': False,
    'pdfplumber': False
}

try:
    from PyPDF2 import PdfReader
    PARSERS_AVAILABLE['pdf'] = True
except ImportError:
    pass

try:
    import pdfplumber
    PARSERS_AVAILABLE['pdfplumber'] = True
except ImportError:
    pass

try:
    import openpyxl
    PARSERS_AVAILABLE['excel'] = True
except ImportError:
    pass

try:
    from docx import Document
    PARSERS_AVAILABLE['word'] = True
except ImportError:
    pass


# File extensions to index
CODE_EXTENSIONS = {
    '.tal', '.cbl', '.cob', '.cobol', '.pli', '.pco', '.cpy',  # Legacy
    '.java', '.py', '.js', '.ts', '.jsx', '.tsx',               # Modern
    '.c', '.cpp', '.h', '.hpp', '.cs', '.go', '.rs', '.rb',     # Systems
    '.sql', '.json', '.yaml', '.yml', '.xml',                    # Data
    '.txt', '.md', '.csv'                                        # Text
}

DOC_EXTENSIONS = {
    '.pdf': 'pdf',
    '.xlsx': 'excel',
    '.xls': 'excel', 
    '.docx': 'word'
}

# Directories to skip
SKIP_DIRS = {
    'node_modules', '.git', '.svn', '__pycache__', 
    'venv', '.venv', 'env', '.env',
    'dist', 'build', 'target', 'out',
    'generated'  # Skip AstraCode generated output
}


def parse_pdf(filepath):
    """Extract text from PDF file."""
    if PARSERS_AVAILABLE['pdfplumber']:
        # pdfplumber gives better results
        try:
            with pdfplumber.open(filepath) as pdf:
                text = '\n'.join(page.extract_text() or '' for page in pdf.pages)
                return text, {'pages': len(pdf.pages)}
        except Exception as e:
            pass
    
    if PARSERS_AVAILABLE['pdf']:
        try:
            reader = PdfReader(filepath)
            text = '\n'.join(page.extract_text() or '' for page in reader.pages)
            return text, {'pages': len(reader.pages)}
        except Exception as e:
            return None, {'error': str(e)}
    
    return None, {'error': 'No PDF parser installed'}


def parse_excel(filepath):
    """Extract text from Excel file."""
    if not PARSERS_AVAILABLE['excel']:
        return None, {'error': 'openpyxl not installed'}
    
    try:
        wb = openpyxl.load_workbook(filepath, read_only=True, data_only=True)
        sheets = []
        full_text = []
        
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            sheet_text = [f'\n=== Sheet: {sheet_name} ===']
            
            for row in sheet.iter_rows(values_only=True):
                row_text = '\t'.join(str(cell) if cell is not None else '' for cell in row)
                if row_text.strip():
                    sheet_text.append(row_text)
            
            sheets.append({'name': sheet_name, 'rows': len(sheet_text) - 1})
            full_text.extend(sheet_text)
        
        wb.close()
        return '\n'.join(full_text), {'sheets': sheets}
    except Exception as e:
        return None, {'error': str(e)}


def parse_word(filepath):
    """Extract text from Word document."""
    if not PARSERS_AVAILABLE['word']:
        return None, {'error': 'python-docx not installed'}
    
    try:
        doc = Document(filepath)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = '\t'.join(cell.text for cell in row.cells)
                if row_text.strip():
                    paragraphs.append(row_text)
        
        return '\n'.join(paragraphs), {'paragraphs': len(paragraphs)}
    except Exception as e:
        return None, {'error': str(e)}


def parse_document(filepath, doc_type):
    """Parse a document based on its type."""
    if doc_type == 'pdf':
        return parse_pdf(filepath)
    elif doc_type == 'excel':
        return parse_excel(filepath)
    elif doc_type == 'word':
        return parse_word(filepath)
    return None, {'error': 'Unknown document type'}


def read_text_file(filepath):
    """Read a text file with encoding detection."""
    encodings = ['utf-8', 'latin-1', 'cp1252', 'ascii']
    
    for encoding in encodings:
        try:
            with open(filepath, 'r', encoding=encoding) as f:
                return f.read()
        except (UnicodeDecodeError, UnicodeError):
            continue
    
    # Binary fallback
    try:
        with open(filepath, 'rb') as f:
            return f.read().decode('utf-8', errors='replace')
    except Exception:
        return None


def extract_functions(content, ext):
    """Extract function/procedure definitions from code."""
    functions = []
    lines = content.split('\n')
    
    # TAL patterns
    if ext in ['.tal']:
        for i, line in enumerate(lines):
            # PROC name
            match = re.match(r'^\s*(PROC|SUBPROC)\s+(\w+)', line, re.IGNORECASE)
            if match:
                functions.append({
                    'name': match.group(2),
                    'type': match.group(1).upper(),
                    'line': i + 1
                })
    
    # COBOL patterns
    elif ext in ['.cbl', '.cob', '.cobol', '.cpy']:
        for i, line in enumerate(lines):
            # PARAGRAPH or SECTION
            match = re.match(r'^\s{7}(\w[\w-]*)\s*(SECTION)?\.', line)
            if match:
                functions.append({
                    'name': match.group(1),
                    'type': 'SECTION' if match.group(2) else 'PARAGRAPH',
                    'line': i + 1
                })
    
    # Java patterns
    elif ext in ['.java']:
        for i, line in enumerate(lines):
            match = re.match(r'^\s*(public|private|protected)?\s*(static)?\s*\w+\s+(\w+)\s*\(', line)
            if match:
                functions.append({
                    'name': match.group(3),
                    'type': 'method',
                    'line': i + 1
                })
    
    # Python patterns
    elif ext in ['.py']:
        for i, line in enumerate(lines):
            match = re.match(r'^\s*def\s+(\w+)\s*\(', line)
            if match:
                functions.append({
                    'name': match.group(1),
                    'type': 'function',
                    'line': i + 1
                })
            match = re.match(r'^\s*class\s+(\w+)', line)
            if match:
                functions.append({
                    'name': match.group(1),
                    'type': 'class',
                    'line': i + 1
                })
    
    # JavaScript/TypeScript patterns
    elif ext in ['.js', '.ts', '.jsx', '.tsx']:
        for i, line in enumerate(lines):
            # function name() or async function name()
            match = re.match(r'^\s*(async\s+)?function\s+(\w+)', line)
            if match:
                functions.append({
                    'name': match.group(2),
                    'type': 'function',
                    'line': i + 1
                })
            # const name = () => or const name = function
            match = re.match(r'^\s*(const|let|var)\s+(\w+)\s*=\s*(async\s+)?(\(|function)', line)
            if match:
                functions.append({
                    'name': match.group(2),
                    'type': 'function',
                    'line': i + 1
                })
    
    return functions


def build_line_index(content, filepath):
    """Build a searchable line index with context."""
    lines = content.split('\n')
    indexed_lines = []
    
    for i, line in enumerate(lines):
        if line.strip():  # Skip empty lines
            indexed_lines.append({
                'line': i + 1,
                'content': line,
                'lower': line.lower()  # Pre-compute for faster search
            })
    
    return indexed_lines


def scan_directory(dirpath, stats):
    """Recursively scan directory for files to index."""
    files_to_index = []
    dirpath = Path(dirpath)
    
    if not dirpath.exists():
        print(f"  ⚠️  Directory not found: {dirpath}")
        return files_to_index
    
    for root, dirs, files in os.walk(dirpath):
        # Skip excluded directories
        dirs[:] = [d for d in dirs if d not in SKIP_DIRS]
        
        for filename in files:
            filepath = Path(root) / filename
            ext = filepath.suffix.lower()
            
            if ext in CODE_EXTENSIONS:
                files_to_index.append({
                    'path': str(filepath),
                    'type': 'code',
                    'ext': ext
                })
                stats['code_files'] += 1
            
            elif ext in DOC_EXTENSIONS:
                doc_type = DOC_EXTENSIONS[ext]
                files_to_index.append({
                    'path': str(filepath),
                    'type': 'document',
                    'doc_type': doc_type,
                    'ext': ext
                })
                stats['doc_files'] += 1
    
    return files_to_index


def build_index(folders, output_path, verbose=False):
    """Build the complete index from multiple folders."""
    
    print("\n🔍 AstraCode Index Builder\n")
    print("=" * 50)
    
    # Show parser status
    print("\n📦 Document Parsers:")
    print(f"   PDF:   {'✅ Ready' if PARSERS_AVAILABLE['pdf'] or PARSERS_AVAILABLE['pdfplumber'] else '❌ pip install pypdf2'}")
    print(f"   Excel: {'✅ Ready' if PARSERS_AVAILABLE['excel'] else '❌ pip install openpyxl'}")
    print(f"   Word:  {'✅ Ready' if PARSERS_AVAILABLE['word'] else '❌ pip install python-docx'}")
    
    # Stats
    stats = {
        'code_files': 0,
        'doc_files': 0,
        'indexed_files': 0,
        'total_lines': 0,
        'functions': 0,
        'errors': 0
    }
    
    # Collect all files
    print(f"\n📁 Scanning {len(folders)} folder(s):")
    all_files = []
    for folder in folders:
        print(f"   - {folder}")
        all_files.extend(scan_directory(folder, stats))
    
    print(f"\n📊 Found: {stats['code_files']} code files, {stats['doc_files']} documents")
    
    # Build index
    index = {
        'version': '1.0',
        'created': datetime.now().isoformat(),
        'folders': [str(Path(f).absolute()) for f in folders],
        'stats': stats,
        'files': {},
        'functions': [],
        'search_index': []
    }
    
    print(f"\n⏳ Indexing files...")
    
    for i, file_info in enumerate(all_files):
        filepath = file_info['path']
        rel_path = filepath
        
        # Try to make relative path
        for folder in folders:
            try:
                rel_path = str(Path(filepath).relative_to(folder))
                break
            except ValueError:
                continue
        
        if verbose:
            print(f"   [{i+1}/{len(all_files)}] {rel_path}")
        elif (i + 1) % 50 == 0:
            print(f"   Processed {i+1}/{len(all_files)} files...")
        
        try:
            if file_info['type'] == 'code':
                content = read_text_file(filepath)
                if content is None:
                    stats['errors'] += 1
                    continue
                
                lines = content.split('\n')
                functions = extract_functions(content, file_info['ext'])
                line_index = build_line_index(content, filepath)
                
                index['files'][filepath] = {
                    'path': filepath,
                    'relative': rel_path,
                    'type': 'code',
                    'ext': file_info['ext'],
                    'lines': len(lines),
                    'functions': functions,
                    'content': line_index
                }
                
                stats['total_lines'] += len(lines)
                stats['functions'] += len(functions)
                stats['indexed_files'] += 1
                
                # Add functions to global index
                for func in functions:
                    index['functions'].append({
                        'name': func['name'],
                        'type': func['type'],
                        'file': filepath,
                        'line': func['line']
                    })
            
            elif file_info['type'] == 'document':
                text, metadata = parse_document(filepath, file_info['doc_type'])
                
                if text is None:
                    if verbose:
                        print(f"      ⚠️  Could not parse: {metadata.get('error', 'unknown')}")
                    stats['errors'] += 1
                    continue
                
                lines = text.split('\n')
                line_index = build_line_index(text, filepath)
                
                index['files'][filepath] = {
                    'path': filepath,
                    'relative': rel_path,
                    'type': 'document',
                    'doc_type': file_info['doc_type'],
                    'ext': file_info['ext'],
                    'lines': len(lines),
                    'metadata': metadata,
                    'content': line_index
                }
                
                stats['total_lines'] += len(lines)
                stats['indexed_files'] += 1
        
        except Exception as e:
            if verbose:
                print(f"      ❌ Error: {e}")
            stats['errors'] += 1
    
    # Update stats in index
    index['stats'] = stats
    
    # Save index
    print(f"\n💾 Saving index to: {output_path}")
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(index, f, indent=2)
    
    file_size = os.path.getsize(output_path)
    size_str = f"{file_size / 1024 / 1024:.1f} MB" if file_size > 1024*1024 else f"{file_size / 1024:.1f} KB"
    
    # Summary
    print(f"\n✅ Index built successfully!")
    print(f"\n📊 Summary:")
    print(f"   Files indexed: {stats['indexed_files']}")
    print(f"   Total lines:   {stats['total_lines']:,}")
    print(f"   Functions:     {stats['functions']:,}")
    print(f"   Errors:        {stats['errors']}")
    print(f"   Index size:    {size_str}")
    print(f"\n📍 Index file: {output_path}")
    print(f"\n💡 Open this folder in VS Code and AstraCode will use the index.")
    
    return index


def main():
    parser = argparse.ArgumentParser(
        description='Build AstraCode index from code and documents',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
    python astra-index.py ./my-project
    python astra-index.py ./src ./docs --output my-index.json
    python astra-index.py . --verbose

Install document parsers:
    pip install pypdf2 openpyxl python-docx
        """
    )
    
    parser.add_argument('folders', nargs='+', help='Folders to index')
    parser.add_argument('-o', '--output', default='.astra-index.json',
                        help='Output index file (default: .astra-index.json)')
    parser.add_argument('-v', '--verbose', action='store_true',
                        help='Show detailed progress')
    
    args = parser.parse_args()
    
    # Resolve folder paths
    folders = [str(Path(f).absolute()) for f in args.folders]
    
    # Build index
    build_index(folders, args.output, args.verbose)


if __name__ == '__main__':
    main()
